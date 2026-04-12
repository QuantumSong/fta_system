"""
文档解析器 — 支持 PDF / Word / Excel / TXT / 图片
"""
import os
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))

from abc import ABC, abstractmethod
from typing import Dict, Any, List
from dataclasses import dataclass, field


@dataclass
class Document:
    """文档数据结构"""
    source: str
    text: str
    images: List[str] = field(default_factory=list)
    tables: List[Dict] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)


class DocumentParser(ABC):
    """文档解析器基类"""

    @abstractmethod
    async def parse(self, file_path: str) -> Document:
        """解析文档"""
        pass


class PDFParser(DocumentParser):
    """PDF解析器 — 使用 pdfplumber 提取文本 + 表格"""

    async def parse(self, file_path: str) -> Document:
        try:
            import pdfplumber

            text_parts: List[str] = []
            tables: List[Dict] = []

            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    # 提取文本
                    page_text = page.extract_text() or ""
                    text_parts.append(page_text)

                    # 提取表格
                    for j, table in enumerate(page.extract_tables() or []):
                        header = table[0] if table else []
                        rows = table[1:] if len(table) > 1 else []
                        tables.append({
                            "page": i + 1,
                            "table_index": j,
                            "header": header,
                            "rows": rows,
                        })
                        # 表格内容也拼入纯文本方便后续分块
                        for row in table:
                            text_parts.append(" | ".join(str(c) if c else "" for c in row))

            return Document(
                source=file_path,
                text="\n".join(text_parts),
                tables=tables,
                metadata={"page_count": len(text_parts), "file_type": "pdf"},
            )
        except Exception as e:
            print(f"PDF解析失败: {e}")
            return Document(source=file_path, text="")


class WordParser(DocumentParser):
    """Word文档解析器"""

    async def parse(self, file_path: str) -> Document:
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)
            text_parts: List[str] = []
            tables: List[Dict] = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # 提取表格
            for i, table in enumerate(doc.tables):
                header = [cell.text for cell in table.rows[0].cells] if table.rows else []
                rows = []
                for row in table.rows[1:]:
                    rows.append([cell.text for cell in row.cells])
                tables.append({"table_index": i, "header": header, "rows": rows})
                for row in table.rows:
                    text_parts.append(" | ".join(cell.text for cell in row.cells))

            return Document(
                source=file_path,
                text="\n".join(text_parts),
                tables=tables,
                metadata={"paragraph_count": len(doc.paragraphs), "file_type": "docx"},
            )
        except Exception as e:
            print(f"Word解析失败: {e}")
            return Document(source=file_path, text="")


class ExcelParser(DocumentParser):
    """Excel解析器"""

    async def parse(self, file_path: str) -> Document:
        try:
            from openpyxl import load_workbook

            wb = load_workbook(file_path, data_only=True)
            text_parts: List[str] = []
            tables: List[Dict] = []

            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheet_data: List[List[str]] = []

                for row in sheet.iter_rows():
                    row_data = [str(cell.value) if cell.value is not None else "" for cell in row]
                    sheet_data.append(row_data)
                    text_parts.append(" | ".join(row_data))

                header = sheet_data[0] if sheet_data else []
                rows = sheet_data[1:] if len(sheet_data) > 1 else []
                tables.append({"sheet_name": sheet_name, "header": header, "rows": rows})

            return Document(
                source=file_path,
                text="\n".join(text_parts),
                tables=tables,
                metadata={"sheet_count": len(wb.sheetnames), "file_type": "xlsx"},
            )
        except Exception as e:
            print(f"Excel解析失败: {e}")
            return Document(source=file_path, text="")


class TxtParser(DocumentParser):
    """纯文本解析器"""

    async def parse(self, file_path: str) -> Document:
        try:
            for enc in ("utf-8", "gbk", "gb2312", "latin-1"):
                try:
                    with open(file_path, "r", encoding=enc) as f:
                        text = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            else:
                text = ""

            return Document(
                source=file_path,
                text=text,
                metadata={"file_type": "txt", "char_count": len(text)},
            )
        except Exception as e:
            print(f"TXT解析失败: {e}")
            return Document(source=file_path, text="")


class ImageParser(DocumentParser):
    """图片解析器 — 返回图片路径，由 LLM 做后续理解"""

    async def parse(self, file_path: str) -> Document:
        try:
            from PIL import Image

            img = Image.open(file_path)
            w, h = img.size
            return Document(
                source=file_path,
                text=f"[图片文件: {os.path.basename(file_path)}, 尺寸: {w}x{h}]",
                images=[file_path],
                metadata={"file_type": "image", "width": w, "height": h, "format": img.format},
            )
        except Exception as e:
            print(f"图片解析失败: {e}")
            return Document(source=file_path, text="")


class DocumentParserFactory:
    """文档解析器工厂"""

    PARSERS = {
        "pdf": PDFParser,
        "docx": WordParser,
        "xlsx": ExcelParser,
        "xls": ExcelParser,
        "txt": TxtParser,
        "png": ImageParser,
        "jpg": ImageParser,
        "jpeg": ImageParser,
    }

    @classmethod
    def get_parser(cls, doc_type: str) -> DocumentParser:
        """获取解析器"""
        parser_class = cls.PARSERS.get(doc_type.lower())
        if parser_class:
            return parser_class()
        # 未知格式当成纯文本处理
        return TxtParser()
